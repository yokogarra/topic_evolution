#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TEN Step 4 — 生成“母题演化研究报告”（Markdown / 可选 Docx / 可选 LLM润色）
=================================================================
输入（来自 Step2 & Step3）：
  - nodes_metrics.csv   （Step3 输出，含 degree/pagerank/community 等）
  - edges_clustered.csv （Step2 输出，母题级跨年边）
  - communities.csv     （Step3 输出，社群概览，可包含 LLM 命名/摘要）

输出：
  - report.md   —— 主报告（Markdown）
  - figures/    —— 可选：为每个社群生成一张年度热度折线图（PNG）
  - report.docx —— 可选：若安装 python-docx 且传 --docx

特色：
  - 自动挑选“核心社群”（按 pagerank / size / 跨年跨度综合排序）
  - 逐社群生成：概览表、年度热度曲线、关键年度、代表主题（来自 label_llm/top_terms）
  - 可选 LLM：将每个社群的说明扩展为 1-2 段精炼文字（--llm 强化）

安装：
  pip install pandas numpy matplotlib python-docx openai

示例：
  python ten_step4_report.py \
    --nodes ./ten_out_step3/nodes_metrics.csv \
    --edges ./ten_out_step2/edges_clustered.csv \
    --comms ./ten_out_step3/communities.csv \
    --outdir ./ten_out_step4 \
    --lang zh --topn 10 --with-fig

  # 可选：LLM 润色社群段落（OpenAI 兼容）
  python ten_step4_report.py \
    --nodes ./ten_out_step3/nodes_metrics.csv \
    --edges ./ten_out_step2/edges_clustered.csv \
    --comms ./ten_out_step3/communities.csv \
    --outdir ./ten_out_step4 \
    --lang zh --topn 10 --with-fig \
    --llm --api openai --model gpt-4o-mini --api-key-env OPENAI_API_KEY
"""
from __future__ import annotations
import os, sys, json, argparse
from typing import List, Dict
import pandas as pd
import numpy as np

# LLM（可选）
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Matplotlib（可选）
try:
    import matplotlib.pyplot as plt
except Exception:
    plt = None

# Docx（可选）
try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None


# ------------------------- 工具函数 -------------------------

def safe_int(x):
    try:
        return int(x)
    except Exception:
        return None


def spark_counts_by_year(df: pd.DataFrame) -> pd.DataFrame:
    """按 year 统计 count 之和（用于折线图/热度分析）。"""
    if 'count' not in df.columns:
        df['count'] = 1
    g = df.groupby('year', as_index=False)['count'].sum().sort_values('year')
    return g


def select_top_communities(nodes: pd.DataFrame, topn: int = 10) -> List[int]:
    # 依据：pagerank（0-1）、size（标准化）、跨度（year_max-year_min+1 标准化）综合
    meta = nodes.copy()
    meta['year'] = pd.to_numeric(meta['year'], errors='coerce')
    meta = meta.dropna(subset=['year'])
    agg = meta.groupby('community').agg(
        size=('id','count'),
        pr=('pagerank','sum'),
        y_min=('year','min'),
        y_max=('year','max')
    ).reset_index()
    agg['span'] = agg['y_max'] - agg['y_min'] + 1
    # 标准化
    for col in ['size','pr','span']:
        v = agg[col].astype(float)
        if v.max() > v.min():
            agg[col+'_z'] = (v - v.min()) / (v.max() - v.min())
        else:
            agg[col+'_z'] = 0.0
    agg['score'] = 0.5*agg['pr_z'] + 0.3*agg['size_z'] + 0.2*agg['span_z']
    top = agg.sort_values('score', ascending=False).head(topn)
    return top['community'].astype(int).tolist()


def llm_expand_paragraph(api: str, api_base: str|None, model: str, api_key_env: str,
                         lang: str, seed_text: str) -> str:
    if api != 'openai' or OpenAI is None:
        return seed_text
    key = os.getenv(api_key_env or 'OPENAI_API_KEY')
    if not key:
        return seed_text
    client = OpenAI(api_key=key, base_url=api_base) if api_base else OpenAI(api_key=key)
    if lang == 'ja':
        prompt = f"以下の素材から、研究報告の段落（3〜5文）を簡潔に作成してください。重複を避け、誇張しない。\n素材:\n{seed_text}"
    elif lang == 'en':
        prompt = f"Write a concise research paragraph (3-5 sentences) from the following notes, avoid repetition or exaggeration.\nNotes:\n{seed_text}"
    else:
        prompt = f"根据以下要点，写一段3-5句的研究性文字，保持克制与客观，避免重复：\n{seed_text}"
    try:
        rsp = client.chat.completions.create(
            model=model,
            messages=[{"role":"user","content":prompt}],
            temperature=0.3, max_tokens=400
        )
        return rsp.choices[0].message.content.strip()
    except Exception:
        return seed_text


# --------------------------- 主流程 ---------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--nodes', required=True, help='nodes_metrics.csv')
    ap.add_argument('--edges', required=True, help='edges_clustered.csv')
    ap.add_argument('--comms', required=True, help='communities.csv')
    ap.add_argument('--outdir', required=True)
    ap.add_argument('--lang', default='zh', choices=['zh','ja','en'])
    ap.add_argument('--topn', type=int, default=8)
    ap.add_argument('--with-fig', action='store_true')
    ap.add_argument('--docx', action='store_true')
    # LLM（可选）
    ap.add_argument('--llm', action='store_true', help='是否使用 LLM 对社群段落进行润色扩写')
    ap.add_argument('--api', default='openai')
    ap.add_argument('--api-base', default=None)
    ap.add_argument('--model', default='gpt-4o-mini')
    ap.add_argument('--api-key-env', default='OPENAI_API_KEY')
    args = ap.parse_args()

    os.makedirs(args.outdir, exist_ok=True)
    figdir = os.path.join(args.outdir, 'figures')
    if args.with_fig and plt is not None:
        os.makedirs(figdir, exist_ok=True)

    n = pd.read_csv(args.nodes)
    e = pd.read_csv(args.edges)
    c = pd.read_csv(args.comms)
    # 列名规范
    for df in (n, e, c):
        df.columns = [x.strip().lower() for x in df.columns]

    # 选择核心社群
    top_cids = select_top_communities(n, topn=args.topn)

    # 开始组装 Markdown 文本
    lines: List[str] = []
    title = {
        'zh': '# 母题演化研究报告',
        'ja': '# 母題進化リポート',
        'en': '# Motif Evolution Report'
    }[args.lang]
    lines.append(title)
    lines.append('')

    # 全局统计
    years = sorted(set(int(y) for y in pd.to_numeric(n['year'], errors='coerce').dropna().astype(int).tolist()))
    lines.append(f"**时间范围**：{years[0]}–{years[-1]}  ")
    lines.append(f"**母题簇数量**：{n['community'].nunique()}  **节点数**：{len(n)}  **边数**：{len(e)}")
    lines.append('')

    # 每个社群
    for rank, cid in enumerate(top_cids, 1):
        sub = n[n['community'] == cid].copy()
        if sub.empty:
            continue
        sub['year'] = pd.to_numeric(sub['year'], errors='coerce')
        sub = sub.dropna(subset=['year'])
        y_min, y_max = int(sub['year'].min()), int(sub['year'].max())
        size = len(sub)
        # 标签与关键词
        labs = [s for s in sub.get('label_llm', '').astype(str).tolist() if isinstance(s, str) and s.strip()]
        key_label = c[c['community']==cid]['label_llm_comm'].fillna('').astype(str).head(1).tolist()
        comm_label = key_label[0] if key_label and key_label[0].strip() else (labs[0] if labs else f'社区 {cid}')
        # 关键词 Top 10
        toks = []
        for t in sub['top_terms'].fillna('').tolist():
            for x in str(t).replace('、',';').replace(',',';').split(';'):
                x = x.strip()
                if x:
                    toks.append(x)
        key_terms = ', '.join(pd.Series(toks).value_counts().head(10).index.tolist())
        # 指标概览
        pr = float(sub['pagerank'].sum())
        deg = int(sub['degree'].sum())
        bmax = float(sub['betweenness'].max())

        lines.append(f"\n## {rank}. {comm_label}  ")
        lines.append(f"**年份**：{y_min}–{y_max}  **规模**：{size}  **PR总和**：{pr:.4f}  **度数总和**：{deg}  **最大桥梁性**：{bmax:.4f}")
        lines.append(f"**关键词**：{key_terms}")

        # 年度热度图
        if args.with_fig and plt is not None:
            ycnt = spark_counts_by_year(sub)
            fig_path = os.path.join(figdir, f"comm_{cid}.png")
            plt.figure()
            plt.plot(ycnt['year'], ycnt['count'], marker='o')
            plt.xlabel('Year'); plt.ylabel('Count'); plt.title(f'Community {cid} — Activity by Year')
            plt.tight_layout(); plt.savefig(fig_path, dpi=150); plt.close()
            rel_path = os.path.relpath(fig_path, args.outdir).replace("\\", "/")
            lines.append(f"![社区 {cid} 年度热度]({rel_path})")

        # 文本说明
        seed = c[c['community']==cid][['label_llm_comm','summary_llm_comm','top_terms_agg']].head(1)
        base_txt = ''
        if not seed.empty:
            lab = seed['label_llm_comm'].astype(str).values[0]
            summ= seed['summary_llm_comm'].astype(str).values[0]
            agg = seed['top_terms_agg'].astype(str).values[0]
            base_txt = '\n'.join([lab, summ, agg])
        else:
            base_txt = key_terms
        if args.llm:
            para = llm_expand_paragraph(args.api, args.api_base, args.model, args.api_key_env, args.lang, base_txt)
        else:
            # 规则生成一小段
            if args.lang == 'ja':
                para = f"本社群は {y_min} 年から {y_max} 年にかけて持続的に現れ、代表的な語は {key_terms} である。"
            elif args.lang == 'en':
                para = f"This community persists from {y_min} to {y_max}, featuring motifs such as {key_terms}."
            else:
                para = f"该社群在 {y_min}–{y_max} 年间持续出现，代表性母题包括：{key_terms}。"
        lines.append("")
        lines.append(para)
        lines.append("")

    # 保存 Markdown
    md_path = os.path.join(args.outdir, 'report.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print('[Saved]', md_path)

    # 可选 Docx
    if args.docx and Document is not None:
        doc = Document()
        doc.add_heading(title, level=0)
        for block in '\n'.join(lines[2:]).split('\n\n'):
            if block.startswith('## '):
                doc.add_heading(block[3:], level=1)
            else:
                p = doc.add_paragraph(block)
                for run in p.runs:
                    run.font.size = Pt(11)
        docx_path = os.path.join(args.outdir, 'report.docx')
        doc.save(docx_path)
        print('[Saved]', docx_path)

    print('\n下一步：可以将 report.md 作为论文附录或分析草稿，或在前端配合 sankey_clustered.html 交互查看。')

if __name__ == '__main__':
    main()
